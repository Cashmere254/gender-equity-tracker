# apps/reports/exporter.py

"""
Server-side DOCX export using python-docx.
Converts a saved Report object and its sections into a
downloadable .docx file.
Called by ReportExportDocxView --- returns an io.BytesIO buffer.
"""

import io
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def export_report_to_docx(report) -> io.BytesIO:
    """
    Builds a Word document from a Report object.
    Returns an io.BytesIO buffer ready for HTTP response.
    """
    doc = DocxDocument()

    # ── Cover page ─────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(report.title)
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x4B, 0x2E, 0x83)  # Purple

    if report.program:
        prog_para = doc.add_paragraph(report.program.name)
        prog_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if report.period_start and report.period_end:
        period = f'{report.period_start} --- {report.period_end}'
        p_para = doc.add_paragraph(period)
        p_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ── Sections ───────────────────────────────────────────────────────────
    for section in report.sections.order_by('position_no'):
        # Section heading
        heading = doc.add_heading(section.heading, level=1)
        heading.runs[0].font.color.rgb = RGBColor(0x4B, 0x2E, 0x83)

        # Section content
        content = section.content_json.get('content', '')
        if content:
            doc.add_paragraph(content)

        doc.add_paragraph('')  # spacer

    # ── Save to buffer ─────────────────────────────────────────────────────
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer